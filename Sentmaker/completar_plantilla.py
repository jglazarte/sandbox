import re
import os
import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
try:
    from docx import Document
except ImportError:
    messagebox.showerror("Error", "La librería python-docx no está instalada. Ejecute el script generar.sh en lugar de este archivo directamente.")
    exit(1)

def iter_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def main():
    plantilla_path = "plantilla.docx"
    
    # Check if plantilla exists. Allow "plantilla.docx"
    if not os.path.exists(plantilla_path):
        messagebox.showerror("Error", f"No se encontró el archivo '{plantilla_path}' en el directorio actual. Por favor, asegúrese de que el archivo existe.")
        return

    try:
        doc = Document(plantilla_path)
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo abrir '{plantilla_path}'. Asegúrese de que es un archivo Word válido y no está corrupto.\n\nError: {e}")
        return

    # Find all unique variables enclosed in brackets
    pattern = r"\[(.*?)\]"
    variables = []
    
    for p in iter_paragraphs(doc):
        matches = re.findall(pattern, p.text)
        for m in matches:
            if m not in variables:
                variables.append(m)

    if not variables:
        messagebox.showinfo("Info", "No se encontraron variables enceradas entre [] en la plantilla.")
        return

    # Create the main window to prompt for variables
    root = tk.Tk()
    root.title("Completar Plantilla de Sentencia")
    root.geometry("800x600")
    
    # Configure grid weights to allow expansion
    root.columnconfigure(0, weight=1)
    root.rowconfigure(0, weight=1)

    # Use ttk theme if available
    style = ttk.Style(root)
    if "clam" in style.theme_names():
        style.theme_use("clam")

    # Main frame with a canvas and scrollbar
    main_frame = ttk.Frame(root)
    main_frame.grid(row=0, column=0, sticky="nsew")
    main_frame.rowconfigure(0, weight=1)
    main_frame.columnconfigure(0, weight=1)

    canvas = tk.Canvas(main_frame, borderwidth=0, highlightthickness=0)
    scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=canvas.yview)
    scrollable_frame = ttk.Frame(canvas)

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )

    canvas_window = canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    
    def on_canvas_configure(e):
        canvas.itemconfig(canvas_window, width=e.width)

    canvas.bind("<Configure>", on_canvas_configure)

    canvas.grid(row=0, column=0, sticky="nsew")
    scrollbar.grid(row=0, column=1, sticky="ns")
    canvas.configure(yscrollcommand=scrollbar.set)

    # Dictionary to hold references to text widgets
    inputs = {}

    ttk.Label(
        scrollable_frame, 
        text="Complete los campos solicitados para generar la sentencia:", 
        font=("Helvetica", 14, "bold")
    ).pack(pady=10, padx=20, anchor="w")

    def focus_next_widget(event):
        event.widget.tk_focusNext().focus()
        return "break"

    def focus_prev_widget(event):
        event.widget.tk_focusPrev().focus()
        return "break"

    for var in variables:
        frame = ttk.Frame(scrollable_frame)
        frame.pack(fill="x", padx=20, pady=5)
        
        lbl = ttk.Label(frame, text=f"[{var}]", font=("Helvetica", 11, "bold"))
        lbl.pack(anchor="w")
        
        text_widget = tk.Text(frame, height=4, font=("Helvetica", 11), wrap="word")
        text_widget.pack(fill="x", pady=5)
        
        text_widget.bind("<Tab>", focus_next_widget)
        text_widget.bind("<Shift-Tab>", focus_prev_widget)
        text_widget.bind("<ISO_Left_Tab>", focus_prev_widget)
        
        inputs[var] = text_widget

    if variables:
        inputs[variables[0]].focus_set()

    def generate_document():
        # Collect answers
        answers = {}
        for var, widget in inputs.items():
            val = widget.get("1.0", tk.END).rstrip('\n')
            answers[var] = val
        
        caratula = answers.get("CARATULA", "Sentencia_Sin_Caratula")
        if not caratula.strip():
            caratula = "Sentencia_Sin_Caratula"
            
        safe_caratula = "".join(c for c in caratula if c.isalnum() or c in (' ', '-', '_')).rstrip()

        # Replace text in the loaded .docx document preserving formatting
        for var, val in answers.items():
            search_text = f"[{var}]"
            
            for p in iter_paragraphs(doc):
                while search_text in p.text:
                    text = p.text
                    start_idx = text.find(search_text)
                    end_idx = start_idx + len(search_text)
                    
                    curr_idx = 0
                    start_run_idx = -1
                    end_run_idx = -1
                    
                    for i, run in enumerate(p.runs):
                        run_len = len(run.text)
                        
                        if start_run_idx == -1 and curr_idx + run_len > start_idx:
                            start_run_idx = i
                            start_char_idx = start_idx - curr_idx
                            
                        if end_run_idx == -1 and curr_idx + run_len >= end_idx:
                            end_run_idx = i
                            end_char_idx = end_idx - curr_idx
                            break
                            
                        curr_idx += run_len
                        
                    start_run = p.runs[start_run_idx]
                    if start_run_idx == end_run_idx:
                        start_run.text = start_run.text[:start_char_idx] + val + start_run.text[end_char_idx:]
                    else:
                        start_run.text = start_run.text[:start_char_idx] + val
                        for i in range(start_run_idx + 1, end_run_idx):
                            p.runs[i].text = ""
                        end_run = p.runs[end_run_idx]
                        end_run.text = end_run.text[end_char_idx:]

        # Save DOCX Document
        try:
            output_filename = f"{safe_caratula}.docx"
            counter = 1
            original_filename = output_filename
            while os.path.exists(output_filename):
                name, ext = os.path.splitext(original_filename)
                output_filename = f"{name}_{counter}{ext}"
                counter += 1

            doc.save(output_filename)
            messagebox.showinfo("Éxito", f"Documento generado exitosamente como:\n{output_filename}")
            root.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Se produjo un error al generar el documento:\n{str(e)}")

    btn_frame = ttk.Frame(root)
    btn_frame.grid(row=1, column=0, sticky="ew", pady=10, padx=20)
    
    btn_generar = ttk.Button(btn_frame, text="Generar Sentencia", command=generate_document)
    btn_generar.pack(side="right")
    
    btn_cancelar = ttk.Button(btn_frame, text="Cancelar", command=root.destroy)
    btn_cancelar.pack(side="right", padx=10)

    # Enable mouse wheel scrolling
    def _on_mousewheel(event):
        canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
    root.bind_all("<MouseWheel>", _on_mousewheel)
    
    # For linux:
    def _on_mousewheel_linux_up(event):
        canvas.yview_scroll(-1, "units")
    def _on_mousewheel_linux_down(event):
        canvas.yview_scroll(1, "units")
        
    root.bind_all("<Button-4>", _on_mousewheel_linux_up)
    root.bind_all("<Button-5>", _on_mousewheel_linux_down)

    root.mainloop()

if __name__ == "__main__":
    main()
